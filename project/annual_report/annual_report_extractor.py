import os

import docx
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker

from project.database import AnnualReport

DB_PATH = r'sqlite:///C:\Users\alsherman\PycharmProjects\annual_report\database\annual_report_heading_only.db'
output_dir_path = r'C:\Users\alsherman\PycharmProjects\annual_report\raw_data\oracle'
company = 'oracle-corporation'

# create object to query database
engine = create_engine(DB_PATH)
Session = sessionmaker(bind=engine)
session = Session()


def main():

    for filename in os.listdir(output_dir_path):
        path = os.path.join(output_dir_path, filename)
        print(path)
        # read docx path
        try:
            doc = docx.Document(path)
        except:
            continue

        # extract text into sections
        sections = {}
        section_name = None  # set as None for first section
        num_paragraphs = len(doc.paragraphs)
        for i in range(0, num_paragraphs):
            p = doc.paragraphs[i]
            # identify new sections (heading font style or lines with all bold text)
            if is_heading(p, heading_types=[1,2]):
                # check for section names that span multiple lines
                # also handle many bold lines in a row (e.g. forms)
                if section_name and new_section == []:
                    section_name = ' '.join([section_name, p.text])
                    continue
                elif section_name:
                    # add previous (completed) section to sections dict
                    sections[section_name] = ' '.join(new_section)
                # start next (or first) section
                section_name = p.text
                new_section = []
            # store all text from each section
            elif section_name:
                new_section.append(p.text)

        # extract metadata from pdf file path
        name = path.split('\\')[-1]
        year = int(name.split('_')[-1].replace('.docx',''))

        for section_name, section in sections.items():

            # create an instance (database row) for a annual report
            annual_report = AnnualReport(
                company=company
                , report_name=name
                , report_year=year
                , section_name=section_name
                , section_text=section['text'].replace('\t',' ').replace('\n',' ')
                , section_type=section['section_type']
            )
            session.add(annual_report)

        # commit (save) all annual reports to the database
        session.commit()


def is_bold(paragraph):
    text_is_bold = []
    for run in paragraph.runs:
        if run.text != '':
            text_is_bold.append(run.bold)

    if text_is_bold != [] and all(text_is_bold):
        return True
    return False


def is_heading(paragraph, heading_types=[1]):
    if isinstance(heading_types, int):
        heading_types = list(heading_types)
    for heading_type in heading_types:
        if 'HEADING {}'.format(heading_type) in paragraph.style.name.upper():
            return True
    return False


if __name__ == "__main__": main()